#!/usr/bin/env python3
"""
Syllabus Parser Agent
Extracts events, assignments, and deadlines from PDF and DOCX syllabi
"""

import re
import logging
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import json

# PDF parsing
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX parsing  
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class SyllabusEvent:
    """Represents an event extracted from a syllabus"""
    
    def __init__(self, title: str, event_type: str, date: datetime, 
                 course: str = "", weight: float = 0.0, location: str = "",
                 description: str = "", confidence: float = 1.0):
        self.title = title
        self.event_type = event_type  # assignment, exam, project, deadline, class
        self.date = date
        self.course = course
        self.weight = weight  # % of final grade
        self.location = location
        self.description = description
        self.confidence = confidence  # 0.0 to 1.0
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            "title": self.title,
            "event_type": self.event_type,
            "date": self.date.isoformat() if self.date else None,
            "course": self.course,
            "weight": self.weight,
            "location": self.location,
            "description": self.description,
            "confidence": self.confidence
        }


class SyllabusParser:
    """Parser for syllabus documents (PDF, DOCX)"""
    
    def __init__(self):
        self.date_patterns = [
            # MM/DD/YYYY or M/D/YYYY
            r'(\d{1,2})/(\d{1,2})/(\d{4})',
            # Month DD, YYYY
            r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{1,2}),?\s+(\d{4})',
            # Mon DD, YYYY (abbreviated)
            r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{1,2}),?\s+(\d{4})',
            # DD Month YYYY
            r'(\d{1,2})\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{4})',
            # YYYY-MM-DD
            r'(\d{4})-(\d{1,2})-(\d{1,2})',
        ]
        
        self.event_keywords = {
            'assignment': ['assignment', 'homework', 'hw', 'problem set', 'ps'],
            'exam': ['exam', 'test', 'quiz', 'midterm', 'final'],
            'project': ['project', 'presentation', 'paper', 'report'],
            'deadline': ['due', 'deadline', 'submit', 'turn in'],
        }
        
        self.weight_pattern = r'(\d+)%'
        self.course_code_pattern = r'[A-Z]{2,4}\s*\d{3,4}'
    
    def parse_pdf(self, pdf_path: str) -> List[SyllabusEvent]:
        """Parse PDF syllabus"""
        if not PDF_AVAILABLE:
            raise RuntimeError("pdfplumber not installed. Run: pip install pdfplumber")
        
        events = []
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                full_text = ""
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
                
                # Extract course info
                course = self._extract_course_code(full_text)
                
                # Extract events
                events = self._extract_events(full_text, course)
                
        except Exception as e:
            logger.error(f"Error parsing PDF {pdf_path}: {e}")
            raise
        
        return events
    
    def parse_docx(self, docx_path: str) -> List[SyllabusEvent]:
        """Parse DOCX syllabus"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        
        events = []
        
        try:
            doc = Document(docx_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Handle tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    full_text += "\n" + row_text
            
            # Extract course info
            course = self._extract_course_code(full_text)
            
            # Extract events
            events = self._extract_events(full_text, course)
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {docx_path}: {e}")
            raise
        
        return events
    
    def _extract_course_code(self, text: str) -> str:
        """Extract course code (e.g., CS 101, MATH 2301)"""
        match = re.search(self.course_code_pattern, text)
        return match.group(0) if match else "Unknown Course"
    
    def _extract_events(self, text: str, course: str) -> List[SyllabusEvent]:
        """Extract events from text"""
        events = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            # Skip empty lines
            if not line.strip():
                continue
            
            # Find dates in line
            dates = self._extract_dates(line)
            
            if dates:
                # Determine event type
                event_type = self._classify_event(line)
                
                # Extract title
                title = self._extract_title(line, event_type)
                
                # Extract weight
                weight = self._extract_weight(line)
                
                # Get context (surrounding lines)
                context_start = max(0, i-1)
                context_end = min(len(lines), i+2)
                context = " ".join(lines[context_start:context_end])
                
                for date in dates:
                    event = SyllabusEvent(
                        title=title,
                        event_type=event_type,
                        date=date,
                        course=course,
                        weight=weight,
                        description=context[:200],
                        confidence=0.8  # Default confidence
                    )
                    events.append(event)
        
        return events
    
    def _extract_dates(self, text: str) -> List[datetime]:
        """Extract dates from text"""
        dates = []
        
        for pattern in self.date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    date = self._parse_date_match(match)
                    if date:
                        dates.append(date)
                except Exception as e:
                    logger.debug(f"Could not parse date from {match.group()}: {e}")
        
        return dates
    
    def _parse_date_match(self, match) -> Optional[datetime]:
        """Parse date from regex match"""
        groups = match.groups()
        text = match.group()
        
        # Try different parsing strategies
        try:
            # MM/DD/YYYY format
            if '/' in text:
                parts = text.split('/')
                if len(parts) == 3:
                    month, day, year = int(parts[0]), int(parts[1]), int(parts[2])
                    return datetime(year, month, day)
            
            # Month name formats
            month_names = {
                'january': 1, 'jan': 1,
                'february': 2, 'feb': 2,
                'march': 3, 'mar': 3,
                'april': 4, 'apr': 4,
                'may': 5,
                'june': 6, 'jun': 6,
                'july': 7, 'jul': 7,
                'august': 8, 'aug': 8,
                'september': 9, 'sep': 9,
                'october': 10, 'oct': 10,
                'november': 11, 'nov': 11,
                'december': 12, 'dec': 12,
            }
            
            text_lower = text.lower()
            for month_name, month_num in month_names.items():
                if month_name in text_lower:
                    # Extract day and year
                    numbers = re.findall(r'\d+', text)
                    if len(numbers) >= 2:
                        day = int(numbers[0])
                        year = int(numbers[1])
                        return datetime(year, month_num, day)
            
            # YYYY-MM-DD format
            if '-' in text:
                parts = text.split('-')
                if len(parts) == 3:
                    year, month, day = int(parts[0]), int(parts[1]), int(parts[2])
                    return datetime(year, month, day)
        
        except Exception as e:
            logger.debug(f"Date parsing error: {e}")
        
        return None
    
    def _classify_event(self, text: str) -> str:
        """Classify event type based on keywords"""
        text_lower = text.lower()
        
        for event_type, keywords in self.event_keywords.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return event_type
        
        return "class"  # Default
    
    def _extract_title(self, text: str, event_type: str) -> str:
        """Extract event title"""
        # Remove date patterns
        cleaned = text
        for pattern in self.date_patterns:
            cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
        
        # Clean up
        cleaned = cleaned.strip()
        cleaned = re.sub(r'\s+', ' ', cleaned)
        
        # Limit length
        if len(cleaned) > 100:
            cleaned = cleaned[:97] + "..."
        
        return cleaned if cleaned else f"{event_type.title()} Event"
    
    def _extract_weight(self, text: str) -> float:
        """Extract percentage weight (e.g., '25%' -> 25.0)"""
        match = re.search(self.weight_pattern, text)
        if match:
            return float(match.group(1))
        return 0.0
    
    def validate_events(self, events: List[SyllabusEvent]) -> Dict[str, Any]:
        """Validate extracted events and flag conflicts/ambiguities"""
        issues = {
            "conflicts": [],
            "ambiguous": [],
            "warnings": []
        }
        
        # Check for date conflicts
        events_by_date = {}
        for event in events:
            date_key = event.date.date() if event.date else None
            if date_key:
                if date_key not in events_by_date:
                    events_by_date[date_key] = []
                events_by_date[date_key].append(event)
        
        # Find conflicts (multiple exams on same day)
        for date, day_events in events_by_date.items():
            exams = [e for e in day_events if e.event_type == 'exam']
            if len(exams) > 1:
                issues["conflicts"].append({
                    "date": date.isoformat(),
                    "events": [e.title for e in exams],
                    "type": "multiple_exams"
                })
        
        # Check for low confidence events
        for event in events:
            if event.confidence < 0.6:
                issues["ambiguous"].append({
                    "title": event.title,
                    "confidence": event.confidence,
                    "date": event.date.isoformat() if event.date else None
                })
        
        # Check for missing weights on graded items
        for event in events:
            if event.event_type in ['assignment', 'exam', 'project'] and event.weight == 0:
                issues["warnings"].append({
                    "title": event.title,
                    "warning": "No grade weight specified"
                })
        
        return issues
    
    def export_events(self, events: List[SyllabusEvent], output_path: str):
        """Export events to JSON"""
        data = {
            "events": [e.to_dict() for e in events],
            "count": len(events),
            "exported_at": datetime.now().isoformat()
        }
        
        with open(output_path, 'w') as f:
            json.dump(data, f, indent=2)
        
        logger.info(f"Exported {len(events)} events to {output_path}")


def test_syllabus_parser():
    """Test syllabus parser"""
    parser = SyllabusParser()
    
    # Test date extraction
    test_text = """
    CS 101 - Introduction to Computer Science
    
    Assignment 1: Due 01/15/2024 (10%)
    Midterm Exam: February 20, 2024 (25%)
    Final Project: May 1, 2024 (30%)
    """
    
    events = parser._extract_events(test_text, "CS 101")
    
    print(f"Extracted {len(events)} events:")
    for event in events:
        print(f"- {event.title} ({event.event_type}) on {event.date} - {event.weight}%")
    
    # Validate
    issues = parser.validate_events(events)
    print(f"\nValidation issues: {issues}")
    
    return len(events) > 0


if __name__ == "__main__":
    success = test_syllabus_parser()
    print(f"\n{'✅ Test passed' if success else '❌ Test failed'}")
