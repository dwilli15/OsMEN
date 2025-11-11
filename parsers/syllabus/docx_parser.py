#!/usr/bin/env python3
"""
DOCX Syllabus Parser

Extracts course information from Word documents (.docx).
Part of v1.4.0 - Syllabus Parser & Calendar Foundation.
"""

import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional
import json

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")


class DOCXSyllabusParser:
    """Parse DOCX syllabi to extract course information and events"""
    
    def __init__(self):
        self.course_info = {}
        self.events = []
        self.assignments = []
        
        # Reuse patterns from PDF parser
        self.date_patterns = [
            r'\b(\d{1,2})/(\d{1,2})/(\d{2,4})\b',
            r'\b(\d{1,2})-(\d{1,2})-(\d{2,4})\b',
            r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{1,2}),?\s+(\d{4})\b',
            r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?\s+(\d{1,2}),?\s+(\d{4})\b',
        ]
        
        self.exam_keywords = ['exam', 'test', 'midterm', 'final', 'quiz']
        self.assignment_keywords = ['assignment', 'homework', 'project', 'paper', 'essay', 'lab', 'due']
    
    def parse_docx(self, docx_path: str) -> Dict[str, Any]:
        """
        Parse DOCX syllabus
        
        Args:
            docx_path: Path to DOCX file
        
        Returns:
            Dictionary with course info, events, and assignments
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        docx_path = Path(docx_path)
        
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX not found: {docx_path}")
        
        # Extract text from DOCX
        text = self._extract_text(docx_path)
        
        # Parse extracted text
        return self._parse_text(text)
    
    def _extract_text(self, docx_path: Path) -> str:
        """Extract text from DOCX document"""
        doc = Document(docx_path)
        text = ""
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract tables (syllabi often have schedules in tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text += row_text + "\n"
        
        return text
    
    def _parse_text(self, text: str) -> Dict[str, Any]:
        """Parse extracted text (same logic as PDF parser)"""
        # Extract course information
        self.course_info = self._extract_course_info(text)
        
        # Extract dates
        dates = self._extract_dates(text)
        
        # Extract events and assignments
        self.events = self._extract_events(text, dates)
        self.assignments = self._extract_assignments(text, dates)
        
        return {
            "course_info": self.course_info,
            "events": self.events,
            "assignments": self.assignments,
            "raw_dates": dates
        }
    
    def _extract_course_info(self, text: str) -> Dict[str, Any]:
        """Extract course information"""
        info = {
            "course_code": None,
            "course_name": None,
            "instructor": None,
            "semester": None,
            "year": None
        }
        
        # Course code
        course_code_match = re.search(r'\b([A-Z]{2,4})\s*(\d{3,4}[A-Z]?)\b', text)
        if course_code_match:
            info["course_code"] = f"{course_code_match.group(1)} {course_code_match.group(2)}"
        
        # Instructor
        instructor_patterns = [
            r'Instructor:\s*([^\n]+)',
            r'Professor:\s*([^\n]+)',
            r'Dr\.\s+([A-Z][a-z]+\s+[A-Z][a-z]+)',
        ]
        for pattern in instructor_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                info["instructor"] = match.group(1).strip()
                break
        
        # Semester
        semester_match = re.search(r'(Fall|Spring|Summer|Winter)\s+(\d{4})', text, re.IGNORECASE)
        if semester_match:
            info["semester"] = semester_match.group(1).capitalize()
            info["year"] = int(semester_match.group(2))
        
        return info
    
    def _extract_dates(self, text: str) -> List[Dict[str, Any]]:
        """Extract all dates from text"""
        dates = []
        
        for pattern in self.date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                date_str = match.group(0)
                date_obj = self._parse_date(match)
                if date_obj:
                    dates.append({
                        "text": date_str,
                        "date": date_obj,
                        "position": match.start()
                    })
        
        dates.sort(key=lambda x: x["position"])
        return dates
    
    def _parse_date(self, match) -> Optional[datetime]:
        """Parse regex match into datetime"""
        try:
            groups = match.groups()
            
            if len(groups) >= 3:
                if groups[0].isdigit():
                    month, day, year = int(groups[0]), int(groups[1]), int(groups[2])
                    if year < 100:
                        year += 2000
                    return datetime(year, month, day)
                else:
                    month_name = groups[0]
                    day = int(groups[1])
                    year = int(groups[2])
                    
                    month_map = {
                        'jan': 1, 'january': 1, 'feb': 2, 'february': 2,
                        'mar': 3, 'march': 3, 'apr': 4, 'april': 4,
                        'may': 5, 'jun': 6, 'june': 6,
                        'jul': 7, 'july': 7, 'aug': 8, 'august': 8,
                        'sep': 9, 'september': 9, 'oct': 10, 'october': 10,
                        'nov': 11, 'november': 11, 'dec': 12, 'december': 12,
                    }
                    
                    month = month_map.get(month_name.lower())
                    if month:
                        return datetime(year, month, day)
        except (ValueError, AttributeError):
            pass
        
        return None
    
    def _extract_events(self, text: str, dates: List[Dict]) -> List[Dict[str, Any]]:
        """Extract events"""
        events = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            for keyword in self.exam_keywords:
                if keyword in line_lower:
                    date = self._find_nearest_date(line, dates, i, lines)
                    if date:
                        events.append({
                            "type": "exam",
                            "title": line.strip()[:100],
                            "date": date.isoformat(),
                            "description": line.strip()
                        })
                    break
        
        return events
    
    def _extract_assignments(self, text: str, dates: List[Dict]) -> List[Dict[str, Any]]:
        """Extract assignments"""
        assignments = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            for keyword in self.assignment_keywords:
                if keyword in line_lower and 'due' in line_lower:
                    date = self._find_nearest_date(line, dates, i, lines)
                    if date:
                        assignments.append({
                            "type": "assignment",
                            "title": line.strip()[:100],
                            "due_date": date.isoformat(),
                            "description": line.strip()
                        })
                    break
        
        return assignments
    
    def _find_nearest_date(self, line: str, dates: List[Dict], 
                          line_idx: int, all_lines: List[str]) -> Optional[datetime]:
        """Find nearest date"""
        for date_info in dates:
            if date_info["text"] in line:
                return date_info["date"]
        
        window_text = ' '.join(all_lines[max(0, line_idx-2):min(len(all_lines), line_idx+3)])
        
        for date_info in dates:
            if date_info["text"] in window_text:
                return date_info["date"]
        
        return None


def main():
    """Test DOCX parser"""
    parser = DOCXSyllabusParser()
    
    print("DOCX Syllabus Parser")
    print("=" * 50)
    
    if not PYTHON_DOCX_AVAILABLE:
        print("❌ python-docx not installed")
        print("   Install with: pip install python-docx")
        return
    
    print("✅ python-docx available")
    
    # Test with sample text
    test_text = """
    MATH 250 - Calculus II
    Spring 2025
    Professor: Dr. Robert Johnson
    
    Important Dates:
    January 20, 2025 - First Class
    March 10, 2025 - Midterm Exam
    May 15, 2025 - Final Exam
    
    Homework 1 due February 3, 2025
    Homework 2 due March 1, 2025
    """
    
    result = parser._parse_text(test_text)
    print("\nParsed Information:")
    print(json.dumps(result, indent=2, default=str))


if __name__ == "__main__":
    main()
